from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = r"C:\Users\Artem\Downloads\diplom\outputs\code_explanation\Фрагмент_кода_ClickerControls_с_объяснением.docx"

CODE_LINES = [
    ("Фрагмент 1. Подписка на VR-события и возврат кликера на стойку", None),
    (46, "private void OnEnable()"),
    (47, "{"),
    (48, "    if (grabInteractable != null)"),
    (49, "    {"),
    (50, "        grabInteractable.activated.AddListener(OnActivated);"),
    (51, "        grabInteractable.selectEntered.AddListener(OnSelectEntered);"),
    (52, "        grabInteractable.selectExited.AddListener(OnSelectExited);"),
    (53, "    }"),
    (54, ""),
    (55, "    EnsureVisibleAndInteractable();"),
    (56, "    if (!isHeld)"),
    (57, "    {"),
    (58, "        DockOnStand();"),
    (59, "    }"),
    (60, "}"),
    ("Фрагмент 2. Команды кликера: слайд, таймер и активация", None),
    (88, "public void NextSlide()"),
    (89, "{"),
    (90, "    if (sessionManager != null) sessionManager.NextSlide();"),
    (91, "}"),
    (93, "public void ToggleTimer()"),
    (94, "{"),
    (95, "    if (sessionManager != null) sessionManager.ToggleTimer();"),
    (96, "}"),
    (108, "private void OnActivated(ActivateEventArgs args)"),
    (109, "{"),
    (110, "    if (useActivateForNextSlide)"),
    (111, "    {"),
    (112, "        NextSlide();"),
    (113, "    }"),
    (114, "}"),
    ("Фрагмент 3. Захват, отпускание и фиксация кликера", None),
    (186, "private void OnSelectEntered(SelectEnterEventArgs args)"),
    (187, "{"),
    (188, "    isHeld = true;"),
    (189, "    if (clickerRigidbody == null) return;"),
    (190, ""),
    (191, "    clickerRigidbody.isKinematic = false;"),
    (192, "    clickerRigidbody.useGravity = false;"),
    (193, "}"),
    (195, "private void OnSelectExited(SelectExitEventArgs args)"),
    (196, "{"),
    (197, "    isHeld = false;"),
    (198, "    DockOnStand();"),
    (199, "}"),
    (201, "private void DockOnStand()"),
    (202, "{"),
    (203, "    if (!dockOnStandWhenReleased || clickerRigidbody == null) return;"),
    (204, ""),
    (205, "    safePosition = standRestPosition;"),
    (206, "    safeRotation = Quaternion.Euler(standRestEulerAngles);"),
    (207, "    clickerRigidbody.isKinematic = false;"),
    (208, "    clickerRigidbody.velocity = Vector3.zero;"),
    (209, "    clickerRigidbody.angularVelocity = Vector3.zero;"),
    (210, "    clickerRigidbody.useGravity = false;"),
    (211, "    clickerRigidbody.isKinematic = true;"),
    (212, "    transform.SetPositionAndRotation(safePosition, safeRotation);"),
    (213, "}"),
]

SPEECH_BLOCKS = [
    (
        "Фрагмент 1. OnEnable: подключение VR-событий",
        "В этом фрагменте я показываю, как кликер подключается к системе VR-взаимодействия Unity. Метод OnEnable вызывается каждый раз, когда объект становится активным. Сначала код проверяет, найден ли компонент grabInteractable. Это компонент XR Interaction Toolkit, который отвечает за возможность взять объект VR-контроллером и получать события взаимодействия. Если компонент существует, скрипт подписывает свои методы на три события: activated, selectEntered и selectExited. Событие activated отвечает за активацию объекта, например нажатие кнопки на контроллере. Событие selectEntered срабатывает, когда пользователь берет кликер. Событие selectExited срабатывает, когда пользователь отпускает кликер. После подписки вызывается EnsureVisibleAndInteractable, чтобы объект был видимым и доступным. Затем проверяется, удерживается ли кликер сейчас. Если он не в руке пользователя, вызывается DockOnStand, и кликер возвращается на стойку. Это нужно, чтобы объект не терялся в сцене и не проваливался через поверхность.",
    ),
    (
        "Фрагмент 2. NextSlide, ToggleTimer и OnActivated: команды тренировки",
        "Во втором фрагменте показано, как кликер передает команды основному сценарию тренировки. Метод NextSlide не переключает слайды самостоятельно, а обращается к sessionManager. Это сделано специально: вся логика тренировки находится в одном центральном менеджере, а кликер только отправляет команду. Такая структура проще для поддержки. Метод ToggleTimer работает аналогично: он проверяет, существует ли sessionManager, и вызывает у него переключение таймера. Метод OnActivated является обработчиком VR-события активации. Он получает аргумент ActivateEventArgs, где Unity хранит данные о событии. Внутри метода есть настройка useActivateForNextSlide. Если она включена, то при активации кликера вызывается NextSlide. То есть пользователь в VR нажимает на контроллере кнопку активации, а приложение реагирует переходом к следующему слайду.",
    ),
    (
        "Фрагмент 3. OnSelectEntered, OnSelectExited и DockOnStand: физика кликера",
        "Третий фрагмент отвечает за поведение кликера в момент захвата и отпускания. Когда пользователь берет кликер VR-контроллером, Unity вызывает OnSelectEntered. Переменная isHeld становится true, то есть скрипт запоминает, что объект сейчас в руке. Затем выполняется проверка Rigidbody. Rigidbody нужен для физического поведения объекта. Если его нет, метод сразу завершается через return, чтобы не получить ошибку. После этого isKinematic переводится в false, чтобы объект мог корректно взаимодействовать с системой захвата, а useGravity отключается, чтобы кликер не падал вниз. Когда пользователь отпускает объект, вызывается OnSelectExited: isHeld становится false, после чего запускается DockOnStand. Метод DockOnStand возвращает кликер в безопасную позицию на стойке. Он сначала проверяет, разрешен ли возврат и есть ли Rigidbody. Затем задает safePosition и safeRotation, обнуляет скорость и вращение, отключает гравитацию, фиксирует объект через isKinematic = true и устанавливает точные координаты и поворот через transform.SetPositionAndRotation. Именно этот блок решает практическую проблему: кликер не должен бесконечно падать или исчезать, он всегда возвращается на понятное место.",
    ),
]

LINE_EXPLANATIONS = [
    ("46", "`private` означает, что метод доступен только внутри этого класса. `void` означает, что метод ничего не возвращает. `OnEnable` - стандартный метод Unity, который вызывается при включении компонента. Скобки `()` показывают, что параметров у метода нет."),
    ("47", "Открывающая фигурная скобка `{` начинает тело метода `OnEnable`. Все строки до соответствующей закрывающей скобки относятся к этому методу."),
    ("48", "`if` задает условие. В скобках проверяется, что поле `grabInteractable` не равно `null`. Знак `!=` означает «не равно». `null` означает отсутствие объекта. Проверка нужна, чтобы не вызвать ошибку, если компонент не найден."),
    ("49", "Открывается блок кода, который выполнится только при истинности условия из строки 48."),
    ("50", "`grabInteractable.activated` - обращение через точку к событию активации VR-объекта. `AddListener(OnActivated)` подписывает метод `OnActivated` на это событие. Точка `.` означает переход к полю или методу объекта. Точка с запятой `;` завершает команду."),
    ("51", "`selectEntered` - событие захвата объекта VR-контроллером. Когда пользователь берет кликер, Unity вызывает событие, а оно запускает метод `OnSelectEntered`. Это связывает действие руки в VR с кодом."),
    ("52", "`selectExited` - событие отпускания объекта. Подписка на `OnSelectExited` нужна, чтобы после отпускания вернуть кликер на стойку."),
    ("53", "Закрывающая фигурная скобка `}` завершает блок условия `if (grabInteractable != null)`."),
    ("54", "Пустая строка не выполняет действий. Она нужна только для визуального разделения кода."),
    ("55", "Вызывается метод `EnsureVisibleAndInteractable`. Он проверяет, что кликер активен, видим, имеет включенные коллайдеры и доступен для взаимодействия."),
    ("56", "`if (!isHeld)` проверяет, что кликер сейчас не удерживается пользователем. Символ `!` означает логическое «не»: условие истинно, когда `isHeld` равно `false`."),
    ("57", "Начинается блок кода, который выполнится, если кликер не находится в руке пользователя."),
    ("58", "Вызывается метод `DockOnStand()`. Он возвращает кликер на заданную позицию стойки. Это решает проблему, когда кликер мог упасть или оказаться не там, где ожидает пользователь."),
    ("59", "Закрывается блок условия из строки 56."),
    ("60", "Закрывается тело метода `OnEnable`."),
    ("88", "`public` означает, что метод можно вызвать извне, например из UI-кнопки на кликере. `NextSlide` - метод перехода к следующему слайду. `void` снова означает отсутствие возвращаемого значения."),
    ("89", "Начинается тело метода `NextSlide`."),
    ("90", "Если `sessionManager` существует, вызывается `sessionManager.NextSlide()`. `sessionManager` - менеджер тренировки, который знает текущий слайд и управляет презентацией. Запись в одну строку допустима, потому что после условия идет одна команда."),
    ("91", "Закрывается метод `NextSlide`."),
    ("93", "Объявляется публичный метод `ToggleTimer`. Его задача - переключать таймер между запуском и паузой."),
    ("94", "Открывается тело метода."),
    ("95", "Если менеджер сессии найден, вызывается `ToggleTimer()` у менеджера. Так кнопка на кликере не управляет таймером напрямую, а передает команду центральному объекту управления."),
    ("96", "Закрывается метод `ToggleTimer`."),
    ("108", "`OnActivated` - обработчик события активации XR-объекта. `ActivateEventArgs args` - параметр с данными события. Тип `ActivateEventArgs` хранит информацию о том, кто и как активировал объект."),
    ("109", "Начинается тело метода `OnActivated`."),
    ("110", "Проверяется настройка `useActivateForNextSlide`. Если она включена, активация кликера будет сразу переключать слайд."),
    ("111", "Открывается блок условия."),
    ("112", "Вызывается метод `NextSlide()`, то есть активация VR-кликера приводит к переходу на следующий слайд."),
    ("113", "Закрывается блок условия."),
    ("114", "Закрывается метод `OnActivated`."),
    ("186", "`OnSelectEntered` - обработчик события захвата. Он вызывается, когда пользователь берет кликер VR-контроллером. Параметр `SelectEnterEventArgs args` содержит сведения о событии захвата."),
    ("187", "Начинается тело метода."),
    ("188", "`isHeld = true;` фиксирует состояние: кликер сейчас удерживается пользователем. Знак `=` здесь означает присваивание значения."),
    ("189", "Если `clickerRigidbody` отсутствует, выполняется `return`, то есть метод досрочно завершается. Это защита от ошибки обращения к несуществующему компоненту."),
    ("190", "Пустая строка отделяет проверку безопасности от дальнейших действий с физикой."),
    ("191", "`clickerRigidbody.isKinematic = false;` включает физическое поведение объекта. Когда `isKinematic` равно `false`, объект может участвовать в физике Unity."),
    ("192", "`useGravity = false` отключает гравитацию. Это важно, чтобы кликер не падал вниз при удержании и не проваливался через поверхность."),
    ("193", "Закрывается метод `OnSelectEntered`."),
    ("195", "`OnSelectExited` - обработчик отпускания VR-объекта. Он вызывается, когда пользователь отпускает кликер."),
    ("196", "Открывается тело метода."),
    ("197", "`isHeld = false;` означает, что кликер больше не находится в руке пользователя."),
    ("198", "Вызывается `DockOnStand()`, чтобы вернуть кликер на стойку после отпускания."),
    ("199", "Закрывается метод `OnSelectExited`."),
    ("201", "`DockOnStand` - вспомогательный метод возврата кликера на фиксированную позицию. Он нужен для стабильности VR-сцены."),
    ("202", "Начинается тело метода."),
    ("203", "Проверяются два условия. `!dockOnStandWhenReleased` означает, что возврат на стойку отключен. `||` означает логическое «или». Если возврат отключен или нет Rigidbody, метод завершается через `return`."),
    ("204", "Пустая строка визуально отделяет проверку от основного алгоритма возврата."),
    ("205", "`safePosition = standRestPosition;` записывает в безопасную позицию координаты стойки. Эти координаты заранее заданы в полях скрипта."),
    ("206", "`Quaternion.Euler(standRestEulerAngles)` преобразует углы Эйлера в поворот Unity. Это нужно, чтобы вернуть кликер не только в позицию, но и в правильный угол."),
    ("207", "`isKinematic = false` временно переводит Rigidbody в обычный режим перед сбросом скоростей."),
    ("208", "`velocity = Vector3.zero` обнуляет линейную скорость. `Vector3.zero` означает вектор `(0, 0, 0)`. Это убирает остаточное движение объекта."),
    ("209", "`angularVelocity = Vector3.zero` обнуляет угловую скорость, то есть вращение кликера."),
    ("210", "`useGravity = false` отключает гравитацию, чтобы объект не начал снова падать после возврата."),
    ("211", "`isKinematic = true` фиксирует объект после возврата. В этом состоянии физика не будет сдвигать кликер со стойки."),
    ("212", "`transform.SetPositionAndRotation(safePosition, safeRotation)` одновременно задает позицию и поворот объекта. Через запятую передаются два аргумента: сначала позиция, затем поворот."),
    ("213", "Закрывается метод `DockOnStand`."),
]

SYMBOLS = [
    ("private", "модификатор доступа: метод виден только внутри класса."),
    ("public", "модификатор доступа: метод можно вызвать из других классов или привязать к UI-кнопке."),
    ("void", "тип возвращаемого значения: метод выполняет действие, но ничего не возвращает."),
    ("()", "круглые скобки после имени метода содержат параметры. Если внутри пусто, параметров нет."),
    ("{ }", "фигурные скобки обозначают начало и конец блока кода."),
    (";", "точка с запятой завершает инструкцию C#."),
    (".", "точка означает обращение к полю, свойству, событию или методу объекта."),
    (",", "запятая разделяет аргументы метода, например позицию и поворот."),
    ("=", "оператор присваивания: справа значение, слева переменная или свойство."),
    ("==", "оператор сравнения «равно»."),
    ("!=", "оператор сравнения «не равно»."),
    ("!", "логическое отрицание: превращает true в false и наоборот."),
    ("||", "логическое «или»: условие истинно, если истинна хотя бы одна часть."),
    ("if", "условный оператор: выполняет блок только при истинном условии."),
    ("return", "досрочный выход из метода."),
    ("null", "отсутствие ссылки на объект."),
    ("true / false", "логические значения: истина и ложь."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name="Times New Roman", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def add_body(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_code_block(document):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F8F8F8")
    set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""
    for number, code in CODE_LINES:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        if code is None:
            p.paragraph_format.space_before = Pt(4)
            run = p.add_run(str(number))
            set_run_font(run, name="Consolas", size=9, bold=True, color="1F4E5F")
            continue
        num = p.add_run(f"{number:>3} | ")
        set_run_font(num, name="Consolas", size=9, color="808080")
        run = p.add_run(code if code else " ")
        set_run_font(run, name="Consolas", size=9, color="1F2933")
    if cell.paragraphs and not cell.paragraphs[0].text:
        cell._element.remove(cell.paragraphs[0]._element)
    return table


def add_table(document, headers, rows, widths=None, font_size=9):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "1F4E5F")
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=font_size, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=font_size)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    return table


document = Document()
section = document.sections[0]
section.orientation = WD_ORIENT.PORTRAIT
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.2)

normal = document.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(12)

add_heading(document, "Фрагмент кода для защиты", 1)
add_body(
    document,
    "Для защиты выбран фрагмент скрипта ClickerControls.cs. Он связан с VR-взаимодействием, потому что использует компонент XRGrabInteractable: пользователь берет кликер контроллером, отпускает его, активирует действие и через кликер управляет слайдами и таймером. Фрагмент также показывает исправление практической проблемы проекта: кликер возвращается на стойку и не проваливается через поверхность.",
)

add_heading(document, "Код для вставки в PowerPoint", 2)
add_code_block(document)

add_heading(document, "Короткое объяснение для устного ответа", 2)
add_body(
    document,
    "Скрипт подписывается на события VR-кликера: захват, отпускание и активацию. При захвате кликер переводится в состояние удержания, при отпускании возвращается на стойку. Методы NextSlide и ToggleTimer передают команды центральному менеджеру тренировки. Благодаря этому VR-объект не работает сам по себе, а является частью общего сценария выступления.",
)

add_heading(document, "Готовый рассказ по фрагментам", 2)
for title, speech in SPEECH_BLOCKS:
    add_heading(document, title, 2)
    add_body(document, speech)

add_heading(document, "Полное объяснение строк кода", 2)
add_table(document, ["Строка", "Объяснение"], LINE_EXPLANATIONS, widths=[1.6, 16.6], font_size=8.5)

add_heading(document, "Что означают символы и ключевые слова", 2)
add_table(document, ["Элемент", "Значение"], SYMBOLS, widths=[3.2, 15.0], font_size=9)

add_heading(document, "Как отвечать на возможные вопросы комиссии", 2)
qa_rows = [
    ("Почему используется XRGrabInteractable?", "Это готовый компонент XR Interaction Toolkit, который позволяет брать объект VR-контроллером и получать события захвата, отпускания и активации."),
    ("Зачем проверять объект на null?", "Чтобы избежать ошибки NullReferenceException, если ссылка на компонент или менеджер не назначена."),
    ("Почему кликер возвращается на стойку?", "Так пользователь всегда знает, где находится объект. Кроме того, это исправляет проблему падения и проваливания через поверхность."),
    ("Почему команды идут через sessionManager?", "Менеджер сессии хранит общую логику тренировки. Кликер только отправляет команды, а не управляет всей сценой напрямую."),
    ("Что делает isKinematic?", "Когда значение true, Rigidbody фиксируется и физика не сдвигает объект. Когда false, объект может участвовать в физическом взаимодействии."),
]
add_table(document, ["Вопрос", "Ответ"], qa_rows, widths=[5.0, 13.2], font_size=9)

document.save(OUTPUT)
print(OUTPUT)
