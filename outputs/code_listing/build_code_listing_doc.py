from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(r"C:\Users\Artem\Downloads\diplom\PublicSpeakingVR")
OUTPUT = Path(r"C:\Users\Artem\Downloads\diplom\outputs\code_listing\РџСЂРёР»РѕР¶РµРЅРёРµ Р‘ - Р›РёСЃС‚РёРЅРі РєРѕРґР°.docx")

SCRIPT_FILES = sorted((PROJECT_ROOT / "Assets" / "Scripts").glob("*.cs"))

DESCRIPTIONS = {
    "AudienceReactionController.cs": "СѓРїСЂР°РІР»СЏРµС‚ СЂРµР°РєС†РёРµР№ Р°СѓРґРёС‚РѕСЂРёРё, СЃРѕСЃС‚РѕСЏРЅРёСЏРјРё РїРµСЂСЃРѕРЅР°Р¶РµР№ Рё РІРѕСЃРїСЂРѕРёР·РІРµРґРµРЅРёРµРј Р·РІСѓРєРѕРІС‹С… СЂРµР°РєС†РёР№.",
    "ClickerControls.cs": "РѕС‚РІРµС‡Р°РµС‚ Р·Р° СЂР°Р±РѕС‚Сѓ VR-РєР»РёРєРµСЂР°, РµРіРѕ С„РёР·РёС‡РµСЃРєРѕРµ РїРѕРІРµРґРµРЅРёРµ, Р·Р°С…РІР°С‚ Рё РїРµСЂРµРєР»СЋС‡РµРЅРёРµ СЃР»Р°Р№РґРѕРІ.",
    "MainMenuController.cs": "СЃРѕР·РґР°РµС‚ РіР»Р°РІРЅРѕРµ РјРµРЅСЋ РІ СЂРµР¶РёРјРµ World Space, РѕР±СЂР°Р±Р°С‚С‹РІР°РµС‚ РІС‹Р±РѕСЂ РґР»РёС‚РµР»СЊРЅРѕСЃС‚Рё С‚СЂРµРЅРёСЂРѕРІРєРё Рё Р·Р°РїСѓСЃРє СЃС†РµРЅС‹.",
    "PCControlFallback.cs": "РѕР±РµСЃРїРµС‡РёРІР°РµС‚ С‚РµСЃС‚РѕРІРѕРµ СѓРїСЂР°РІР»РµРЅРёРµ СЃ РєР»Р°РІРёР°С‚СѓСЂС‹ Рё РјС‹С€Рё РїСЂРё РѕС‚СЃСѓС‚СЃС‚РІРёРё VR-РіР°СЂРЅРёС‚СѓСЂС‹.",
    "PresentationSessionManager.cs": "РєРѕРѕСЂРґРёРЅРёСЂСѓРµС‚ С‚СЂРµРЅРёСЂРѕРІРѕС‡РЅСѓСЋ СЃРµСЃСЃРёСЋ, С‚Р°Р№РјРµСЂ, СЃР»Р°Р№РґС‹, СЃС‚Р°С‚РёСЃС‚РёРєСѓ Рё РїРµСЂРµС…РѕРґС‹ РјРµР¶РґСѓ СЃРѕСЃС‚РѕСЏРЅРёСЏРјРё.",
    "SlideChanger.cs": "РІС‹РїРѕР»РЅСЏРµС‚ СЃРјРµРЅСѓ СЃР»Р°Р№РґРѕРІ РїСЂРµР·РµРЅС‚Р°С†РёРё Рё РѕС‚СЃР»РµР¶РёРІР°РµС‚ С‚РµРєСѓС‰РёР№ РЅРѕРјРµСЂ СЃР»Р°Р№РґР°.",
    "SpeakerTimer.cs": "СЂРµР°Р»РёР·СѓРµС‚ С‚Р°Р№РјРµСЂ РІС‹СЃС‚СѓРїР»РµРЅРёСЏ, Р·Р°РїСѓСЃРє, РїР°СѓР·Сѓ, Р·Р°РІРµСЂС€РµРЅРёРµ РѕС‚СЃС‡РµС‚Р° Рё РѕС‚РѕР±СЂР°Р¶РµРЅРёРµ РІСЂРµРјРµРЅРё.",
    "VRRuntimeOptimizer.cs": "СЃРѕРґРµСЂР¶РёС‚ РЅР°СЃС‚СЂРѕР№РєРё РѕРїС‚РёРјРёР·Р°С†РёРё VR-СЂРµР¶РёРјР° Рё РїР°СЂР°РјРµС‚СЂРѕРІ РєР°С‡РµСЃС‚РІР° РґР»СЏ Р°РІС‚РѕРЅРѕРјРЅРѕР№ РіР°СЂРЅРёС‚СѓСЂС‹.",
    "VRThumbstickMovement.cs": "РґРѕР±Р°РІР»СЏРµС‚ РїРµСЂРµРјРµС‰РµРЅРёРµ РїРѕР»СЊР·РѕРІР°С‚РµР»СЏ РІ VR РїСЂРё РїРѕРјРѕС‰Рё СЃС‚РёРєР° РєРѕРЅС‚СЂРѕР»Р»РµСЂР°.",
}

KEYWORDS = {
    "abstract", "as", "base", "bool", "break", "byte", "case", "catch", "char", "class", "const",
    "continue", "decimal", "default", "delegate", "do", "double", "else", "enum", "event", "explicit",
    "extern", "false", "finally", "fixed", "float", "for", "foreach", "if", "implicit", "in", "int",
    "interface", "internal", "is", "lock", "long", "namespace", "new", "null", "object", "operator",
    "out", "override", "params", "private", "protected", "public", "readonly", "ref", "return",
    "sbyte", "sealed", "short", "sizeof", "stackalloc", "static", "string", "struct", "switch", "this",
    "throw", "true", "try", "typeof", "uint", "ulong", "unchecked", "unsafe", "ushort", "using",
    "virtual", "void", "volatile", "while", "var", "get", "set", "value",
}

UNITY_TYPES = {
    "MonoBehaviour", "GameObject", "Transform", "Vector2", "Vector3", "Quaternion", "Color", "Camera",
    "Canvas", "TextMeshProUGUI", "Button", "Image", "Rigidbody", "Collider", "AudioClip", "AudioSource",
    "Material", "Renderer", "SerializeField", "Header", "Range", "Tooltip", "SceneManager", "Mathf",
    "Time", "Debug", "PlayerPrefs", "InputActionReference", "CharacterController",
}

TOKEN_PATTERN = re.compile(
    r"(//.*$|\"(?:\\.|[^\"\\])*\"|'(?:\\.|[^'\\])'|\b\d+(?:\.\d+)?f?\b|\b[A-Za-z_][A-Za-z0-9_]*\b)",
    re.MULTILINE,
)


def set_run_font(run, name="Times New Roman", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(document, text, size=14, align=WD_ALIGN_PARAGRAPH.CENTER):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=True)


def add_body(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def add_caption(document, number, file_name, description):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"Р›РёСЃС‚РёРЅРі Р‘.{number} вЂ” РЎРєСЂРёРїС‚ {file_name}")
    set_run_font(run, size=12, bold=True)

    desc = document.add_paragraph()
    desc.paragraph_format.space_after = Pt(4)
    desc.paragraph_format.first_line_indent = Cm(1.25)
    desc_run = desc.add_run(f"РќР°Р·РЅР°С‡РµРЅРёРµ: {description}")
    set_run_font(desc_run, size=11)


def add_code_segment(paragraph, text, color):
    if not text:
        return
    run = paragraph.add_run(text)
    set_run_font(run, name="Consolas", size=14, color=color)


def add_code_line(document, line_number, line):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    number = paragraph.add_run(f"{line_number:>4} | ")
    set_run_font(number, name="Consolas", size=14, color="808080")

    line = line.rstrip("\n").replace("\t", "    ")
    position = 0
    for match in TOKEN_PATTERN.finditer(line):
        if match.start() > position:
            add_code_segment(paragraph, line[position:match.start()], "222222")

        token = match.group(0)
        if token.startswith("//"):
            color = "008000"
        elif token.startswith("\"") or token.startswith("'"):
            color = "A31515"
        elif re.fullmatch(r"\d+(?:\.\d+)?f?", token):
            color = "098658"
        elif token in KEYWORDS:
            color = "0000FF"
        elif token in UNITY_TYPES:
            color = "2B91AF"
        else:
            color = "222222"

        add_code_segment(paragraph, token, color)
        position = match.end()

    if position < len(line):
        add_code_segment(paragraph, line[position:], "222222")

    if not line:
        add_code_segment(paragraph, " ", "222222")


document = Document()
section = document.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.2)

normal = document.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(12)

add_heading(document, "РџР РР›РћР–Р•РќРР• Р‘", size=14)
add_heading(document, "Р›РёСЃС‚РёРЅРі РєРѕРґР° РґРёРїР»РѕРјРЅРѕРіРѕ РїСЂРѕРµРєС‚Р°", size=14)
add_body(
    document,
    "Р’ РїСЂРёР»РѕР¶РµРЅРёРё РїСЂРёРІРµРґРµРЅС‹ Р»РёСЃС‚РёРЅРіРё C#-СЃРєСЂРёРїС‚РѕРІ РёР· РїР°РїРєРё Assets/Scripts, РёСЃРїРѕР»СЊР·СѓРµРјС‹С… РІ VR-С‚СЂРµРЅР°Р¶РµСЂРµ РїСѓР±Р»РёС‡РЅРѕРіРѕ РІС‹СЃС‚СѓРїР»РµРЅРёСЏ. Р”Р»СЏ СѓРґРѕР±СЃС‚РІР° Р°РЅР°Р»РёР·Р° РєР°Р¶РґР°СЏ СЃС‚СЂРѕРєР° РєРѕРґР° СЃРЅР°Р±Р¶РµРЅР° РЅРѕРјРµСЂРѕРј, Р° РїРµСЂРµРґ Р»РёСЃС‚РёРЅРіРѕРј СѓРєР°Р·Р°РЅРѕ РЅР°Р·РЅР°С‡РµРЅРёРµ СЃРѕРѕС‚РІРµС‚СЃС‚РІСѓСЋС‰РµРіРѕ СЃРєСЂРёРїС‚Р°.",
)

add_heading(document, "Р‘.1 РЎРєСЂРёРїС‚С‹ VR-С‚СЂРµРЅР°Р¶РµСЂР°", size=13, align=WD_ALIGN_PARAGRAPH.LEFT)

for index, script_path in enumerate(SCRIPT_FILES, start=1):
    file_name = script_path.name
    description = DESCRIPTIONS.get(file_name, "СЃР»СѓР¶РµР±РЅС‹Р№ СЃРєСЂРёРїС‚ РїСЂРѕРµРєС‚Р° Unity.")
    add_caption(document, index, file_name, description)

    lines = script_path.read_text(encoding="utf-8-sig", errors="replace").splitlines()
    for line_number, line in enumerate(lines, start=1):
        add_code_line(document, line_number, line)

document.save(OUTPUT)
print(OUTPUT)

