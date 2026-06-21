from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\Artem\Downloads\diplom\outputs\svetlichny_testing\Светличный с замечаниями - исправлен пункт 2.3.docx")
OUTPUT = Path(r"C:\Users\Artem\Downloads\diplom\outputs\svetlichny_testing\Светличный с замечаниями - исправлены пункты 1.3 и 2.3.docx")


def paragraph_text(element):
    return "".join(node.text or "" for node in element.xpath(".//w:t"))


def set_run_font(run, size=12, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def insert_paragraph_after(cursor, text="", style=None, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
    new_p = OxmlElement("w:p")
    cursor.addnext(new_p)
    paragraph = Paragraph(new_p, cursor.getparent())
    if style:
        paragraph.style = style
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=12, bold=bold)
    return paragraph._p, paragraph


def set_cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def add_caption(cursor, text):
    cursor, p = insert_paragraph_after(cursor, text, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    return cursor


def move_table_after(document, cursor, rows, cols):
    table = document.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    cursor.addnext(tbl)
    return tbl, table


document = Document(SOURCE)
body = document._body._element
children = list(body)

start = None
end = None
for idx, child in enumerate(children):
    if child.tag == qn("w:p") and paragraph_text(child).strip() == "1.3 Обзор и обоснование выбора технологий и инструментов":
        start = idx
    elif start is not None and child.tag == qn("w:p") and paragraph_text(child).strip() == "2 ПРАКТИЧЕСКАЯ ЧАСТЬ":
        end = idx
        break

if start is None or end is None:
    raise RuntimeError("Не удалось найти границы раздела 1.3")

heading = children[start]
for child in children[start + 1:end]:
    body.remove(child)

cursor = heading

paragraphs_before_table = [
    "Выбор технологий и инструментов разработки является важным этапом дипломного проекта, так как от него зависят скорость реализации, удобство отладки, производительность приложения и возможность дальнейшего сопровождения. Для мобильной игры Tap and Catch рассматривались средства, которые позволяют создать приложение для Android, реализовать обработку касаний, игровой цикл, графический вывод и хранение результата пользователя.",
    "В качестве возможных языков программирования были рассмотрены Kotlin, Java и C#. Java традиционно применяется в Android-разработке и обладает большим количеством учебных материалов, однако код на Java обычно получается более объемным. C# удобен при использовании Unity, но для небольшой 2D-игры подключение игрового движка увеличивает сложность проекта. Kotlin был выбран как современный официальный язык Android-разработки, позволяющий писать компактный и читаемый код.",
    "При выборе среды разработки рассматривались Android Studio, IntelliJ IDEA и Unity. Android Studio является основной средой для разработки Android-приложений, содержит инструменты сборки Gradle, эмулятор, редактор ресурсов, средства отладки и профилирования. IntelliJ IDEA удобна для Kotlin-проектов, но требует дополнительной настройки Android-инструментов. Unity подходит для более сложных игр, однако для Tap and Catch его возможности являются избыточными.",
    "Для реализации графики рассматривались XML-интерфейсы Android, Jetpack Compose, SurfaceView с Canvas и игровой движок Unity. XML-разметка хорошо подходит для статичных экранов, но не предназначена для постоянной перерисовки игрового поля. Jetpack Compose удобен для современных интерфейсов, однако для учебной 2D-аркады проще и нагляднее использовать SurfaceView и Canvas, где игровой цикл, отрисовка объектов и обработка столкновений описываются напрямую.",
]

for text in paragraphs_before_table:
    cursor, _ = insert_paragraph_after(cursor, text)

cursor = add_caption(cursor, "Таблица 3 — Сравнение программных средств для реализации приложения")

headers = ["Группа средств", "Рассмотренные варианты", "Преимущества", "Ограничения", "Выбранное решение"]
rows = [
    ["Язык программирования", "Kotlin, Java, C#", "Kotlin и Java имеют прямую поддержку Android; C# удобен в Unity", "Java более многословна; C# требует игрового движка", "Kotlin"],
    ["Среда разработки", "Android Studio, IntelliJ IDEA, Unity", "Android Studio содержит эмулятор, Gradle, отладчик и инструменты Android SDK", "Unity избыточен для простой 2D-игры; IntelliJ IDEA требует дополнительной настройки", "Android Studio"],
    ["Графический вывод", "XML-разметка, Jetpack Compose, SurfaceView Canvas, Unity", "SurfaceView Canvas позволяет напрямую управлять игровым циклом и перерисовкой", "XML и Compose менее удобны для динамичной игровой сцены", "SurfaceView и Canvas"],
    ["Хранение данных", "SharedPreferences, SQLite, Room", "SharedPreferences прост для хранения небольших параметров", "SQLite и Room рассчитаны на более сложные структуры данных", "SharedPreferences"],
    ["Тестирование и запуск", "Android-эмулятор, физическое устройство", "Эмулятор доступен в Android Studio и удобен для демонстрации", "Физическое устройство не всегда доступно при защите", "Android Emulator"],
]

cursor, table = move_table_after(document, cursor, rows=len(rows) + 1, cols=len(headers))
for c, header in enumerate(headers):
    set_cell_text(table.rows[0].cells[c], header, bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9EAF7")

for r, row in enumerate(rows, start=1):
    for c, value in enumerate(row):
        set_cell_text(table.rows[r].cells[c], value, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)

cursor = table._tbl

paragraphs_after_table = [
    "Сравнение показывает, что для дипломного проекта целесообразно использовать не универсальный игровой движок, а стандартный стек Android-разработки. Это связано с тем, что приложение имеет компактную структуру, не требует сложной физики, 3D-графики, сетевого режима или большого количества сцен. Основная задача проекта состоит в демонстрации мобильной разработки, обработки пользовательского ввода и организации простого игрового цикла.",
    "Итоговый технологический стек проекта включает Kotlin, Android Studio, Android SDK, SurfaceView, Canvas и SharedPreferences. Kotlin используется для реализации логики экранов, игрового цикла, подсчета очков и обработки касаний. Android Studio обеспечивает сборку, запуск и отладку приложения. SurfaceView и Canvas отвечают за вывод игрового поля, корзины и падающих объектов. SharedPreferences применяется для хранения лучшего результата пользователя.",
    "Выбранный стек является обоснованным, поскольку соответствует масштабу проекта, не перегружает приложение лишними зависимостями и позволяет показать основные навыки Android-разработки. Кроме того, такой подход облегчает проверку проекта на защите: исходный код, ресурсы интерфейса, логика игры и результаты тестирования находятся в одной среде разработки и могут быть быстро продемонстрированы преподавателю.",
]

for text in paragraphs_after_table:
    cursor, _ = insert_paragraph_after(cursor, text)

document.save(OUTPUT)
print(OUTPUT)
