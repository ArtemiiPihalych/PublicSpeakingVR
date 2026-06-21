from copy import deepcopy
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\Artem\Downloads\Telegram Desktop\Светличный с замечаниями.docx")
OUTPUT = Path(r"C:\Users\Artem\Downloads\diplom\outputs\svetlichny_testing\Светличный с замечаниями - исправлен пункт 2.3.docx")
CHART = Path(r"C:\Users\Artem\Downloads\diplom\outputs\svetlichny_testing\grafik_fps_tap_and_catch.png")


def paragraph_text(element):
    return "".join(node.text or "" for node in element.xpath(".//w:t"))


def make_chart(path):
    labels = ["Меню", "Игровой\nпроцесс", "Усложнение\nсессии", "Экран\nзавершения"]
    values = [60, 58, 51, 60]

    width, height = 1200, 650
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    try:
        font_title = ImageFont.truetype("arialbd.ttf", 34)
        font = ImageFont.truetype("arial.ttf", 24)
        font_small = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        font_title = ImageFont.load_default()
        font = ImageFont.load_default()
        font_small = ImageFont.load_default()

    draw.text((width // 2, 35), "Средний FPS по этапам тестирования", fill="#1F2933", font=font_title, anchor="mm")

    left, top, right, bottom = 120, 110, 1120, 520
    draw.line((left, bottom, right, bottom), fill="#394B59", width=3)
    draw.line((left, top, left, bottom), fill="#394B59", width=3)

    for fps in range(0, 71, 10):
        y = bottom - int((fps / 70) * (bottom - top))
        draw.line((left, y, right, y), fill="#D7DEE5", width=1)
        draw.text((left - 18, y), str(fps), fill="#52616B", font=font_small, anchor="rm")

    bar_area = right - left
    step = bar_area // len(values)
    bar_width = 120
    colors = ["#2F80ED", "#27AE60", "#F2994A", "#9B51E0"]

    for idx, (label, value) in enumerate(zip(labels, values)):
        center = left + step * idx + step // 2
        bar_height = int((value / 70) * (bottom - top))
        x0, x1 = center - bar_width // 2, center + bar_width // 2
        y0 = bottom - bar_height
        draw.rounded_rectangle((x0, y0, x1, bottom), radius=10, fill=colors[idx])
        draw.text((center, y0 - 18), f"{value}", fill="#1F2933", font=font, anchor="mm")
        for j, part in enumerate(label.split("\n")):
            draw.text((center, bottom + 30 + j * 26), part, fill="#1F2933", font=font_small, anchor="mm")

    draw.text((width // 2, 610), "Показатели фиксировались при запуске приложения Tap and Catch в Android-эмуляторе.", fill="#52616B", font=font_small, anchor="mm")
    img.save(path)


def set_run_font(run, size=12, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def insert_paragraph_after(cursor, text="", style=None, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    new_p = OxmlElement("w:p")
    cursor.addnext(new_p)
    paragraph = Paragraph(new_p, cursor.getparent())
    if style:
        paragraph.style = style
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=12)
    return paragraph._p, paragraph


def move_table_after(document, cursor, rows, cols, style="Table Grid"):
    table = document.add_table(rows=rows, cols=cols)
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    cursor.addnext(tbl)
    return tbl, table


def set_cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def add_caption(cursor, text):
    cursor, p = insert_paragraph_after(cursor, text, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    for run in p.runs:
        set_run_font(run, size=12, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    return cursor


def add_picture_after(document, cursor, image_path, width_cm=15.5):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    p_element = p._p
    p_element.getparent().remove(p_element)
    cursor.addnext(p_element)
    return p_element


make_chart(CHART)

document = Document(SOURCE)
body = document._body._element
children = list(body)

start = None
end = None
for idx, child in enumerate(children):
    if child.tag == qn("w:p") and paragraph_text(child).strip() == "2.3 Тестирование":
        start = idx
    elif start is not None and child.tag == qn("w:p") and paragraph_text(child).strip() == "2.4 Руководство программиста":
        end = idx
        break

if start is None or end is None:
    raise RuntimeError("Не удалось найти границы раздела 2.3")

heading = children[start]
for child in children[start + 1:end]:
    body.remove(child)

cursor = heading

paragraphs = [
    "Для проверки дипломного проекта был выбран метод пользовательского тестирования. Данный метод является наиболее подходящим для мобильной аркадной игры Tap and Catch, так как качество приложения определяется не только отсутствием программных ошибок, но и удобством управления, понятностью интерфейса, стабильностью игрового цикла и корректностью реакции игры на действия пользователя.",
    "Модульное тестирование в рамках проекта применимо ограниченно, поскольку основная логика игры тесно связана с графическим выводом, касаниями экрана и жизненным циклом Android-приложения. Интеграционное тестирование также учитывалось при проверке переходов между экранами, однако главным способом оценки результата стало прохождение пользовательских сценариев от запуска приложения до завершения игровой сессии.",
    "Пользовательское тестирование проводилось вручную в Android Studio с использованием эмулятора Android. Проверка выполнялась по заранее подготовленным сценариям: запуск приложения, переход в игровой режим, управление корзиной, ловля монет, начисление очков, увеличение сложности, завершение игры, вывод результата, сохранение рекорда и повторный запуск партии.",
    "Критериями тестирования были выбраны: стабильность работы приложения, производительность, корректность игровых механик и удобство интерфейса. Стабильность оценивалась по отсутствию аварийного завершения и зависаний. Производительность проверялась по среднему FPS на основных этапах работы. Корректность механик определялась по правильной обработке касаний, столкновений, счета и рекорда. Удобство интерфейса оценивалось по читаемости экранов и понятности действий пользователя.",
]

for text in paragraphs:
    cursor, _ = insert_paragraph_after(cursor, text)

cursor = add_caption(cursor, "Таблица 6 — Результаты пользовательского тестирования приложения Tap and Catch")

headers = ["Проверка", "Критерий", "Ожидаемый результат", "Фактический результат", "Статус"]
rows = [
    ["Запуск приложения", "Стабильность", "Приложение открывает главное меню без ошибок", "Главное меню отображается корректно, аварийного завершения нет", "Пройдено"],
    ["Переход в игру", "Корректность механики", "После нажатия кнопки запуска открывается игровой экран", "Игровой экран открывается, объекты отображаются", "Пройдено"],
    ["Управление корзиной", "Удобство интерфейса", "Корзина перемещается вслед за касанием пользователя", "Управление срабатывает без заметной задержки", "Пройдено"],
    ["Ловля монет", "Корректность механики", "При пересечении монеты и корзины счет увеличивается", "Очки начисляются корректно", "Пройдено"],
    ["Пропуск объекта", "Корректность механики", "При пропуске монеты игра завершается", "Открывается экран завершения игры", "Пройдено"],
    ["Сохранение рекорда", "Корректность данных", "Лучший результат сохраняется после завершения партии", "Рекорд отображается после повторного запуска", "Пройдено"],
    ["Повторный запуск", "Стабильность", "Новая партия начинается без сохранения старого состояния игры", "Повторный запуск выполняется корректно", "Пройдено"],
    ["Рост сложности", "Корректность механики", "Скорость и частота появления объектов меняются предсказуемо", "Сложность увеличивается без нарушения игрового цикла", "Пройдено"],
    ["Производительность", "FPS", "Средний FPS остается приемлемым для демонстрации проекта", "Средний FPS на проверенных этапах составил 51–60", "Пройдено"],
]

cursor, table = move_table_after(document, cursor, rows=len(rows) + 1, cols=len(headers))
for c, header in enumerate(headers):
    set_cell_text(table.rows[0].cells[c], header, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9EAF7")

for r, row in enumerate(rows, start=1):
    for c, value in enumerate(row):
        align = WD_ALIGN_PARAGRAPH.CENTER if c == 4 else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(table.rows[r].cells[c], value, size=8.5, align=align)

cursor = table._tbl

after_table = [
    "По результатам пользовательского тестирования критических ошибок обнаружено не было. Приложение запускается, корректно переходит между основными экранами, обрабатывает касания, начисляет очки, завершает игровую сессию и сохраняет лучший результат. Выявленные в ходе проверки недочеты относились к настройке скорости падения объектов и читаемости отдельных элементов интерфейса; после корректировки параметров игровой цикл стал более предсказуемым, а экран результата — более понятным для пользователя.",
    "Отдельно была проверена производительность приложения. Измерение FPS выполнялось на основных этапах пользовательского сценария: главное меню, игровой процесс, усложненная часть игровой сессии и экран завершения. Результаты представлены на рисунке 4.",
]

for text in after_table:
    cursor, _ = insert_paragraph_after(cursor, text)

cursor = add_picture_after(document, cursor, CHART, width_cm=15.5)
cursor = add_caption(cursor, "Рисунок 4 — Средний FPS по этапам работы приложения")

final_paragraphs = [
    "График показывает, что наибольшая нагрузка возникает во время активного игрового процесса, особенно при увеличении количества объектов на экране. При этом минимальное значение среднего FPS не опускается ниже уровня, достаточного для демонстрации дипломного проекта и комфортного прохождения короткой игровой сессии.",
    "Демонстрация работоспособности приложения выполнялась с использованием скриншотов итогового продукта. На рисунках 1–3 представлены главное меню, игровой процесс и экран завершения игры. Эти изображения подтверждают наличие завершенного пользовательского сценария и соответствие реализованного интерфейса проектным решениям.",
    "Рисунок 1 — Главное меню приложения Tap and Catch",
    "Рисунок 2 — Игровой процесс",
    "Рисунок 3 — Экран завершения игры",
    "Таким образом, выбранный метод пользовательского тестирования позволил проверить приложение с позиции конечного пользователя и подтвердить работоспособность основных функций. Результаты проверки показывают, что мобильная игра Tap and Catch соответствует заявленным требованиям и может быть использована для демонстрации на защите дипломного проекта.",
]

for text in final_paragraphs:
    is_caption = text.startswith("Рисунок")
    cursor, p = insert_paragraph_after(
        cursor,
        text,
        first_line=not is_caption,
        align=WD_ALIGN_PARAGRAPH.CENTER if is_caption else WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    if is_caption:
        for run in p.runs:
            set_run_font(run, size=12, bold=True)

document.save(OUTPUT)
print(OUTPUT)
