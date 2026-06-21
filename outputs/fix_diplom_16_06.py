from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


DOC_PATH = Path(r"C:\Users\Artem\Downloads\diplom\ДипломФАА v3.1 16.06.docx")
BACKUP_PATH = Path(r"C:\Users\Artem\Downloads\diplom\ДипломФАА v3.1 16.06.backup.docx")


TASKS = [
    "Проанализировать предметную область и существующие аналоги.",
    "Рассмотреть основные теоретические аспекты разработки VR-приложения.",
    "Выполнить обзор и обосновать выбор технологий и инструментов разработки.",
    "Спроектировать структуру, пользовательские сценарии и интерфейс VR-тренажера.",
    "Реализовать основные программные модули проекта.",
    "Провести тестирование разработанного VR-приложения.",
    "Разработать руководство программиста.",
    "Разработать руководство пользователя.",
    "Выполнить экономическое обоснование проекта.",
    "Рассмотреть требования техники безопасности и охраны труда.",
]

GOST_SOURCES = [
    "ГОСТ 34.601-90. Информационные технологии. Автоматизированные системы. Стадии создания. - Москва: Стандартинформ, 2020. - 12 с.",
    "ГОСТ Р ИСО/МЭК 12207-2010. Информационная технология. Системная и программная инженерия. Процессы жизненного цикла программных средств. - Москва: Стандартинформ, 2012. - 106 с.",
    "ГОСТ 19.504-79. Единая система программной документации. Руководство программиста. Требования к содержанию и оформлению. - Москва: Стандартинформ, 2010. - 7 с.",
    "ГОСТ Р 59795-2021. Информационные технологии. Комплекс стандартов на автоматизированные системы. Автоматизированные системы. Руководство пользователя. Требования к содержанию, оформлению и обозначению. - Москва: Российский институт стандартизации, 2021. - 16 с.",
]


def paragraph_text(element):
    return "".join(node.text or "" for node in element.xpath(".//w:t"))


def set_times(paragraph, size=12):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)


def insert_paragraph_after(cursor, text, template_pPr=None):
    new_p = OxmlElement("w:p")
    if template_pPr is not None:
        new_p.append(deepcopy(template_pPr))
    cursor.addnext(new_p)
    paragraph = Paragraph(new_p, cursor.getparent())
    paragraph.add_run(text)
    set_times(paragraph)
    return paragraph._p, paragraph


def replace_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)
    set_times(paragraph)


if not BACKUP_PATH.exists():
    BACKUP_PATH.write_bytes(DOC_PATH.read_bytes())

document = Document(DOC_PATH)
body = document._body._element

# 1. Replace task paragraphs in the introduction.
paragraphs = document.paragraphs
task_intro_index = None
for i, paragraph in enumerate(paragraphs):
    if paragraph.text.strip().startswith("Для достижения цели были поставлены задачи"):
        task_intro_index = i
        break

if task_intro_index is None:
    raise RuntimeError("Не найден абзац с задачами")

task_start = task_intro_index + 1
task_end = task_start
while task_end < len(paragraphs) and not paragraphs[task_end].text.strip().startswith("Объектом разработки"):
    task_end += 1

template_task = paragraphs[task_start]
template_task_pPr = template_task._p.pPr

for paragraph in paragraphs[task_start:task_end]:
    body.remove(paragraph._p)

cursor = paragraphs[task_intro_index]._p
for index, task in enumerate(TASKS, start=1):
    cursor, paragraph = insert_paragraph_after(cursor, f"{index}. {task}", template_pPr=template_task_pPr)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(1.25)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

document.save(DOC_PATH)

# 2. Reorder and extend source list: GOSTs first.
document = Document(DOC_PATH)
body = document._body._element
paragraphs = document.paragraphs

source_heading_index = None
appendix_index = None
for i, paragraph in enumerate(paragraphs):
    text = paragraph.text.strip()
    if text == "СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ":
        source_heading_index = i
    elif source_heading_index is not None and text == "ПРИЛОЖЕНИЯ":
        appendix_index = i
        break

if source_heading_index is None or appendix_index is None:
    raise RuntimeError("Не найдены границы списка источников")

source_paragraphs = paragraphs[source_heading_index + 1:appendix_index]
source_texts = [p.text.strip() for p in source_paragraphs if p.text.strip()]
source_template = source_paragraphs[0]
source_pPr = source_template._p.pPr

non_gost_sources = []
for text in source_texts:
    if not text.startswith("ГОСТ 34.601-90") and not text.startswith("ГОСТ Р ИСО/МЭК 12207-2010"):
        non_gost_sources.append(text)

for p in source_paragraphs:
    body.remove(p._p)

cursor = paragraphs[source_heading_index]._p
for text in GOST_SOURCES + non_gost_sources:
    cursor, paragraph = insert_paragraph_after(cursor, text, template_pPr=source_pPr)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(0)

document.save(DOC_PATH)

# 3. Update numeric references in the body after source reordering.
document = Document(DOC_PATH)

mapping = {26: 1, 27: 2}
for n in range(1, 26):
    mapping[n] = n + 4
for n in range(28, 80):
    mapping[n] = n + 2


def replace_reference(match):
    raw = match.group(1)
    parts = [part.strip() for part in raw.split(",")]
    if not all(part.isdigit() for part in parts):
        return match.group(0)
    updated = [str(mapping.get(int(part), int(part))) for part in parts]
    return "[" + ", ".join(updated) + "]"


for paragraph in document.paragraphs:
    if paragraph.text.strip() == "СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ":
        break
    new_text = re.sub(r"\[([0-9]+(?:\s*,\s*[0-9]+)*)\]", replace_reference, paragraph.text)
    if new_text != paragraph.text:
        replace_paragraph_text(paragraph, new_text)

document.save(DOC_PATH)
print(DOC_PATH)
